"""Python helpers for parsing and exporting resume files (PDF, DOCX, plain text)."""

import base64
import os
import tempfile


def parse_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    import io
    reader = PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text.strip()


def parse_docx(file_bytes: bytes) -> str:
    from docx import Document
    import io
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def parse_resume_file(file_name: str, file_data_b64: str) -> str:
    """Decode base64 file and extract plain text based on file type."""
    file_bytes = base64.b64decode(file_data_b64)
    ext = os.path.splitext(file_name)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes)
    else:
        return file_bytes.decode("utf-8", errors="replace")


def generate_resume_docx(parsed: dict) -> str:
    """Generate a clean, ATS-friendly DOCX from parsed resume data. Returns base64."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    doc = Document()

    # Remove default styles noise and set margins
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Remove default paragraph spacing
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    def set_cell_border(cell, **kwargs):
        pass  # unused but kept for clarity

    def add_rule():
        """Add a thin horizontal rule paragraph."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_section_heading(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(title.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # Underline the heading with a bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '6633CC')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_bullet(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run('• ' + text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        return p

    # ── Name ──────────────────────────────────────────────────────────────────
    name = parsed.get('name', 'Resume')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(name)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True

    # ── Contact ───────────────────────────────────────────────────────────────
    contact = parsed.get('contact', {})
    contact_parts = [v for k, v in contact.items() if v]
    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run('  |  '.join(contact_parts))
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Summary ───────────────────────────────────────────────────────────────
    if parsed.get('summary'):
        add_section_heading('Summary')
        p = doc.add_paragraph(parsed['summary'])
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.name = 'Calibri'

    # ── Experience ────────────────────────────────────────────────────────────
    experience = parsed.get('experience', [])
    if experience:
        add_section_heading('Experience')
        for i, exp in enumerate(experience):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4 if i > 0 else 2)
            p.paragraph_format.space_after = Pt(0)
            # Role (bold) + Company
            run = p.add_run(exp.get('role', ''))
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            company = exp.get('company', '')
            loc = exp.get('location', '')
            company_str = company + (f'  ·  {loc}' if loc else '')
            if company_str:
                run2 = p.add_run(f'  |  {company_str}')
                run2.font.size = Pt(10)
                run2.font.name = 'Calibri'
                run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            # Dates (right-aligned via tab)
            if exp.get('dates'):
                run3 = p.add_run(f'\t{exp["dates"]}')
                run3.font.size = Pt(9)
                run3.font.name = 'Calibri'
                run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                # Add right tab stop
                from docx.oxml import OxmlElement
                pPr = p._p.get_or_add_pPr()
                tabs = OxmlElement('w:tabs')
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '9360')  # ~6.5 inches
                tabs.append(tab)
                pPr.append(tabs)
            for bullet in exp.get('bullets', []):
                add_bullet(bullet)

    # ── Education ─────────────────────────────────────────────────────────────
    education = parsed.get('education', [])
    if education:
        add_section_heading('Education')
        for i, edu in enumerate(education):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4 if i > 0 else 2)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(edu.get('institution', ''))
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if edu.get('dates'):
                run2 = p.add_run(f'\t{edu["dates"]}')
                run2.font.size = Pt(9)
                run2.font.name = 'Calibri'
                run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                pPr = p._p.get_or_add_pPr()
                tabs = OxmlElement('w:tabs')
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '9360')
                tabs.append(tab)
                pPr.append(tabs)
            degree_str = edu.get('degree', '')
            if edu.get('gpa'):
                degree_str += f'  ·  GPA {edu["gpa"]}'
            if degree_str:
                p2 = doc.add_paragraph(degree_str)
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(1)
                if p2.runs:
                    p2.runs[0].font.size = Pt(10)
                    p2.runs[0].font.name = 'Calibri'
                    p2.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = parsed.get('skills', [])
    if skills:
        add_section_heading('Skills')
        p = doc.add_paragraph(', '.join(skills))
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.name = 'Calibri'

    # ── Projects ──────────────────────────────────────────────────────────────
    projects = parsed.get('projects', [])
    if projects:
        add_section_heading('Projects')
        for i, proj in enumerate(projects):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4 if i > 0 else 2)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(proj.get('name', ''))
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if proj.get('dates'):
                run2 = p.add_run(f'\t{proj["dates"]}')
                run2.font.size = Pt(9)
                run2.font.name = 'Calibri'
                run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                pPr = p._p.get_or_add_pPr()
                tabs = OxmlElement('w:tabs')
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '9360')
                tabs.append(tab)
                pPr.append(tabs)
            techs = proj.get('technologies', [])
            if techs:
                p2 = doc.add_paragraph('Technologies: ' + ', '.join(techs))
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(1)
                if p2.runs:
                    p2.runs[0].font.size = Pt(9)
                    p2.runs[0].font.name = 'Calibri'
                    p2.runs[0].italic = True
                    p2.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            for bullet in proj.get('bullets', []):
                add_bullet(bullet)

    # ── Certifications ────────────────────────────────────────────────────────
    certs = parsed.get('certifications', [])
    if certs:
        add_section_heading('Certifications')
        for cert in certs:
            add_bullet(cert)

    # ── Serialize ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode('utf-8')
